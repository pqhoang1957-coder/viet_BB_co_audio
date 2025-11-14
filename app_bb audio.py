# =========================
# app_bien_ban_streamlit_fix.py
# =========================
import streamlit as st
from google import genai
from docx import Document
import os

# =========================
# 1) Nhập GEMINI_API_KEY
# =========================
API_KEY = st.text_input("Nhập GEMINI_API_KEY:", type="password")
if not API_KEY:
    st.warning("Vui lòng nhập GEMINI_API_KEY để tiếp tục.")
    st.stop()

client = genai.Client(api_key=API_KEY)

# =========================
# 2) Giao diện upload file hoặc dán văn bản
# =========================
st.title("🤖 Trợ Lý Viết Biên Bản (VBI HCM - Gemini)")

uploaded_file = st.file_uploader(
    "Tải lên file ghi âm (.mp3, .wav, .flac):",
    type=["mp3", "wav", "flac"]
)

meeting_notes = st.text_area(
    "HOẶC Dán Nội Dung Cuộc Họp Thô vào đây:",
    height=200
)

# =========================
# 3) Xử lý khi nhấn nút "Soạn thảo biên bản"
# =========================
if st.button("Soạn thảo biên bản"):

    if uploaded_file is None and not meeting_notes.strip():
        st.warning("Vui lòng tải lên file hoặc dán nội dung cuộc họp.")
        st.stop()

    with st.spinner("Đang xử lý..."):

        system_instruction = """
        Bạn là 1 chuyên gia trong lĩnh vực tạo báo cáo buổi họp của công ty bảo hiểm phi nhân thọ VBI Hồ Chí MInh với hơn 10 năm kinh nghiệm. Chatbot hỗ trợ soạn thảo báo cáo từ các ghi chú hoặc từ nội dung do người dùng cung cấp. Báo cáo được trình bày rõ ràng, chính xác có cấu trúc chuẩn gồm: thời gian họp, địa điểm họp, thành phần tham dự, nội dung chính của buổi họp, các quyết định, yêu cầu, hành động tiếp theo và người phụ trách thực hiện. Chatbot có thể viết biên bản bằng tiếng Việt hoặc tiếng Anh tuỳ theo yêu cầu. Chatbot giữ văn phong trang trọng, ngắn gọn và chính xác. Nếu thông tin chưa đầy đủ, chưa rõ, Chatbot sẽ chủ động hỏi lại để làm rõ trước khi soạn báo cáo.
        Chatbot cũng hỗ trợ người dùng chuyển báo cáo sang các định dạng trình bày khác nhau, ví dụ: email tóm tắt, văn bản hành chính.
        Nhiệm vụ 1: Phân tích và tổ chức thông tin đầu vào
        - Xác định và phân loại thông tin chính từ nội dung thô.
        - Nhận diện các yếu tố cốt lõi: thời gian, địa điểm, đối tượng.
        - Phân chia nội dung thành: thảo luận, vấn đề nổi bật, ý kiến đóng góp, quyết định.
        - Các thông tin được cung cấp có thể rời rạc nhưng phại tập hợp lại thành cùng đoạn văn bản nếu có cùng nội dung, cùng chủ đề.
        Nhiệm vụ 2: Soạn thảo báo cáo họp theo định dạng chuẩn
        - Gồm: Tiêu đề, thời gian, địa điểm, người tham dự, nội dung, kết luận, hành động tiếp theo.
        - Sử dụng ngôn ngữ trang trọng, mạch lạc, hành chính, rõ ràng.
        - Đảm bảo ngữ pháp, chính tả và định dạng thống nhất.
        Nhiệm vụ 3: Tùy chỉnh định dạng báo cáo theo yêu cầu
        - Chuyển báo cáo thành email, văn bản chính thức hoặc bản để trình bày.
        - Điều chỉnh văn phong theo đối tượng người nhận.
        - Tùy biến độ chi tiết theo yêu cầu.
        Nhiệm vụ 4: Rà soát và tối ưu báo cáo
        - Kiểm tra lỗi chính tả, ngữ pháp và logic tổng thể.
        - Gợi ý cải thiện nội dung chưa rõ ràng.
        - Đảm bảo thông tin không bị trùng lặp, mâu thuẫn.
        Quy tắc hoạt động:
        1. Chỉ sử dụng thông tin đã được xác minh từ người dùng, không tự suy luận, không bịa số liệu.
        2. Luôn hỏi lại nếu thông tin chưa rõ ràng hoặc thiếu, cần thiết yêu cầu gửi biểu số liệu để phân tích. Các từ viết tắt chưa rõ phải hỏi và ghi nhớ cho lần sau
        3. Văn phong hành chính, trang trọng, ngắn gọn.
        4. Tôn trọng yêu cầu về gửi định dạng của người dùng.
        5. Không xuất nội dung dưới dạng tệp hoặc mẫu định sẵn.
        6. Đảm bảo tính logic, mạch lạc trong toàn bộ văn bản.
        7. Giữ tính riêng tư và bảo mật nội dung cuộc họp.
        """

        gem_file = None  # Khởi tạo để xóa tạm nếu có

        try:
            # --- Ưu tiên: file audio ---
            if uploaded_file is not None:
                st.info("Đang upload file audio lên Gemini...")

                # Lưu file tạm trên server Streamlit
                tmp_path = f"/tmp/{uploaded_file.name}"
                with open(tmp_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())

                # Upload lên Gemini
                gem_file = client.files.upload(file=tmp_path)

                contents = [
                    system_instruction,
                    gem_file,
                    "Hãy phiên âm file và viết biên bản họp hoàn chỉnh."
                ]
                model_to_use = "gemini-2.5-pro"

            # --- Nếu chỉ có text ---
            else:
                contents = system_instruction + "\n\nNội dung cuộc họp:\n" + meeting_notes
                model_to_use = "gemini-2.5-flash"

            # Gọi API tạo biên bản
            response = client.models.generate_content(
                model=model_to_use,
                contents=contents,
                config={"temperature": 0.1}
            )

            biens_ban_text = response.text

            # Hiển thị biên bản
            st.subheader("📄 Biên bản hoàn chỉnh")
            st.text_area("Kết quả", biens_ban_text, height=300)

            # Xuất ra Word
            doc = Document()
            doc.add_heading('BIÊN BẢN CUỘC HỌP', 0)
            doc.add_paragraph(biens_ban_text)
            word_filename = f"{uploaded_file.name.rsplit('.',1)[0] if uploaded_file else 'BienBan'}_BienBan.docx"
            word_path = f"/tmp/{word_filename}"
            doc.save(word_path)

            st.download_button(
                label="📥 Tải biên bản Word",
                data=open(word_path, "rb").read(),
                file_name=word_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        except Exception as e:
            st.error(f"Đã xảy ra lỗi: {e}")

        finally:
            # Xóa file tạm trên Gemini nếu upload audio
            if gem_file is not None:
                try:
                    client.files.delete(name=gem_file.name)
                    st.success("✅ Đã dọn file tạm trên Gemini.")
                except Exception as e_del:
                    st.warning(f"Không xóa được file tạm trên Gemini: {e_del}")

